import os
from re import A
from turtle import width
from typing import IO, Dict, List, Union
from pypdf import PageObject, PdfReader, PdfWriter
import pypdf
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.colors import black, gray
import io

from docx import Document
from docx.shared import Cm


class PDF_Utils:
    def __init__(
        self,
        folder_path: str,
        pdf_path: str,
        template_path: str,
        cover: str,
        presentation: Union[str, None] = None,
    ) -> None:
        self.folder_path = folder_path
        self.template_path = template_path
        self.pdf_path = pdf_path
        self.cover = cover
        self.presentation = presentation

    def read_pdfs(self) -> Dict[str, PdfReader]:
        files: Dict[str, PdfReader] = {}
        for filename in os.listdir(self.folder_path):
            if "TemplateWord" not in filename:
                pdf_path = os.path.join(self.folder_path, filename)
            if os.path.isfile(pdf_path) and pdf_path.endswith(".pdf"):
                files[pdf_path] = PdfReader(pdf_path)
        return files

    def merge_files(self, file_dict: Dict[str, PdfReader], output_path: str) -> None:
        merger = PdfWriter()
        for file in file_dict.keys():
            merger.append(file)
        with open(output_path, "wb") as f:
            merger.write(f)
        merger.close()

    def make_index_page(self, index: Dict[int, List[Dict]]) -> PageObject:
        packet = io.BytesIO()

        # Create a new PDF with ReportLab
        cv = canvas.Canvas(packet, pagesize=letter)
        cv.setFont("Helvetica", 12)

        for idx, (key, val) in enumerate(index.items()):
            cv.setFillColor(black)
            cv.setFontSize(11)

            next_topic_jump = len(val) * 20 + 10 if len(val) > 0 else idx * 30

            cv.drawString(100, 800 - next_topic_jump, f"Index {key}: {len(val)} topics")
            for indexes, items in enumerate(val):
                cv.setFillColor(gray)
                cv.setFontSize(8)
                cv.drawString(
                    110,
                    790 - indexes * 20 - (idx + 1) * 30,
                    f"- {items['Descrição']}",
                )

        cv.save()  # Finalize the new PDF we just created
        packet.seek(0)  # Go to the beginning of the BytesIO buffer
        new_pdf = PdfReader(packet)

        # Read the template PDF
        template_pdf = PdfReader(self.template_path)
        template_page = template_pdf.pages[0]

        # Create a new PageObject for the output PDF that combines the template and the new index
        output_page = PageObject.create_blank_page(
            width=template_page.mediabox.width,
            height=template_page.mediabox.height,
        )

        # Merge template content with the new index content
        output_page.merge_page(template_page)  # Copy content from the template
        output_page.merge_page(
            new_pdf.pages[0]
        )  # Add the new index without overwriting
        return output_page

    def alt_index_page(self, index: Dict[int, List[Dict]]):
        document = Document()
        for idx, (key, val) in enumerate(index.items()):
            document.add_heading(f"Index {key}: {len(val)} topics", level=1)
            for items in val:
                document.add_paragraph(f"- {items['Descrição']}", style="List Bullet")
        document.save("alt_index.docx")

    def create_index_and_merge(
        self, merged_file_path: str, index: Dict[int, List[Dict]]
    ) -> None:
        merged_file = PdfReader(merged_file_path)
        merged_with_index = PdfWriter()

        # TODO
        # change for adding outlines to the actual pages
        # need to find the pages in the document

        index_page = self.make_index_page(index)
        merged_with_index.add_page(index_page)

        # count the pages on each added document, but how am i gonna track
        # what docs are being added?
        # -e

        section_indexes = [
            "Section: " + str(title) for i, title in enumerate(index.keys())
        ]

        # section_sub_indexes = []

        for i, v in enumerate(index.values()):
            parent_outline = merged_with_index.add_outline_item(
                title=section_indexes[i], page_number=i + 1, bold=True
            )
            for _, values in enumerate(v):
                merged_with_index.add_outline_item(
                    title=values["Descrição"],
                    # change here to right addresses
                    page_number=i + 1,
                    parent=parent_outline,
                )

        merged_with_index.add_page(merged_file.pages[0])

        with open(merged_file_path, "wb") as f:
            merged_with_index.write(f)

        merged_with_index.close()

    def final_doc(self, merged_path: str):
        w = PdfWriter()
        r = PdfReader(merged_path)

        capa = PdfReader(self.cover)
        apresentação = (
            PdfReader(self.presentation) if self.presentation is not None else None
        )
